#!/usr/bin/env python3
"""전문가 에이전트 분석 결과 → MS Word 주간 전략 레포트 생성 + 운영 서버 업로드"""

import sys
import json
import datetime
import requests
from pathlib import Path

PROJECT_ROOT  = Path(__file__).parent.parent
REPORT_DIR    = PROJECT_ROOT / "reports" / "weekly"
ENV_FILE      = PROJECT_ROOT / ".env"
PROD_API_BASE = "https://aimsp-backend.onrender.com"


def load_env() -> dict:
    env = {}
    if ENV_FILE.exists():
        for line in ENV_FILE.read_text().splitlines():
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                k, v = line.split('=', 1)
                env[k.strip()] = v.strip().strip('"').strip("'")
    return env


def upload_to_server(output_path: Path, meta: dict) -> str | None:
    """생성된 Word 파일을 운영 서버에 업로드하여 다운로드 URL 반환 (JSON+Base64)"""
    import base64
    env        = load_env()
    api_token  = env.get("API_SECRET_TOKEN") or os.environ.get("API_SECRET_TOKEN", "")
    week_start = meta.get("weekStart", "")
    week_end   = meta.get("weekEnd", "")
    title      = f"AI MSP 주간 전략 레포트 ({week_start} ~ {week_end})"

    try:
        content_b64 = base64.b64encode(output_path.read_bytes()).decode()
        payload = {
            "title":        title,
            "weekStart":    week_start,
            "weekEnd":      week_end,
            "articleCount": meta.get("articleCount", 0),
            "insightCount": meta.get("insightCount", 0),
            "content":      content_b64,
        }
        headers = {"X-API-Token": api_token} if api_token else {}
        resp = requests.post(
            f"{PROD_API_BASE}/api/weekly-reports/upload",
            json=payload,
            headers=headers,
            timeout=60
        )
        if resp.ok:
            data         = resp.json()
            report_id    = data.get("id")
            download_url = f"{PROD_API_BASE}/api/weekly-reports/{report_id}/download"
            return download_url
        else:
            print(f"  ⚠️  업로드 실패: {resp.status_code} {resp.text[:200]}", file=sys.stderr)
            return None
    except Exception as e:
        print(f"  ⚠️  업로드 오류: {e}", file=sys.stderr)
        return None


def load_json(path: Path) -> dict:
    if not path.exists():
        print(f"  ⚠️  파일 없음: {path}", file=sys.stderr)
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def generate_word(workspace: Path, meta: dict, analyses: list[dict]):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

    COLOR_TITLE   = RGBColor(0x1F, 0x4E, 0x79)
    COLOR_SECTION = RGBColor(0x1F, 0x4E, 0x79)
    COLOR_EXPERT  = RGBColor(0x2E, 0x74, 0xB5)
    COLOR_COMP    = RGBColor(0x70, 0x30, 0xA0)
    COLOR_META    = RGBColor(0x88, 0x88, 0x88)
    COLOR_FOOTER  = RGBColor(0xAA, 0xAA, 0xAA)
    FONT          = "맑은 고딕"

    def p(text, size=11, bold=False, color=None, align=None, indent=0):
        para = doc.add_paragraph()
        if align:
            para.alignment = align
        if indent:
            para.paragraph_format.left_indent = Cm(indent)
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = FONT
        if color:
            run.font.color.rgb = color
        return para

    def bullet(text):
        if not text:
            return
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(str(text))
        run.font.size = Pt(11)
        run.font.name = FONT

    week_start = meta.get("weekStart", "")
    week_end   = meta.get("weekEnd", "")
    art_count  = meta.get("articleCount", 0)
    ins_count  = meta.get("insightCount", 0)

    # ── 제목 ──────────────────────────────────────────
    p(f"AI MSP 주간 전략 레포트",
      size=20, bold=True, color=COLOR_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(f"({week_start} ~ {week_end})",
      size=14, bold=True, color=COLOR_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)

    now = datetime.datetime.now().strftime("%Y년 %m월 %d일 %H시 %M분")
    p(f"생성일시: {now}  |  분석 기사: {art_count}건  |  인사이트: {ins_count}건",
      size=10, color=COLOR_META, align=WD_ALIGN_PARAGRAPH.CENTER)
    p("작성: AI MSP 전문가 · ITO 전문가 · IT 전략 전문가 (Claude Opus 4)",
      size=10, color=COLOR_META, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ── 경영진 요약 ───────────────────────────────────
    it_strategy = next((a for a in analyses if a.get("expertRole") == "IT 전략 전문가"), {})
    exec_summary = it_strategy.get("executiveSummary", "")
    if exec_summary:
        p("■ 경영진 요약", size=13, bold=True, color=COLOR_SECTION)
        p(exec_summary, size=11)
        doc.add_paragraph()

    # ── 섹션 1: 경쟁사 동향 ───────────────────────────
    p("1. 지난주 주요 경쟁사 동향", size=15, bold=True, color=COLOR_SECTION)
    doc.add_paragraph()

    competitors = ["LG CNS", "SK AX", "베스핀글로벌", "PwC Korea"]
    for comp in competitors:
        p(f"▶ {comp}", size=13, bold=True, color=COLOR_COMP)
        for analysis in analyses:
            role = analysis.get("expertRole", "")
            by_comp = analysis.get("competitorAnalysis", {}).get("byCompetitor", [])
            comp_data = next((c for c in by_comp
                              if c.get("name", "") in comp or comp in c.get("name", "")), None)
            if not comp_data:
                continue
            p(f"[{role}]", size=11, bold=True, color=COLOR_EXPERT)
            for mv in comp_data.get("movements", comp_data.get("itoMovements",
                                    comp_data.get("strategicMove", []))):
                if isinstance(mv, list):
                    for m in mv:
                        bullet(m)
                else:
                    bullet(mv)
            impl = comp_data.get("implication", "")
            if impl:
                p(f"시사점: {impl}", size=11, indent=0.5)
        doc.add_paragraph()

    # 경쟁사 동향 총평
    p("[종합 총평]", size=12, bold=True, color=COLOR_EXPERT)
    for analysis in analyses:
        summary = analysis.get("competitorAnalysis", {}).get("summary", "")
        if summary:
            p(f"· {analysis.get('expertRole')}: {summary}", size=11)
    doc.add_paragraph()

    # ── 섹션 2: AI 사업 Trend ─────────────────────────
    p("2. AI 사업 Trend", size=15, bold=True, color=COLOR_SECTION)
    doc.add_paragraph()

    # 트렌드 총평
    for analysis in analyses:
        trend_summary = analysis.get("trendAnalysis", {}).get("summary", "")
        if trend_summary:
            p(f"[{analysis.get('expertRole')}]", size=11, bold=True, color=COLOR_EXPERT)
            p(trend_summary, size=11)

    doc.add_paragraph()

    # 주요 트렌드 상세
    trend_titles_seen = set()
    for analysis in analyses:
        role   = analysis.get("expertRole", "")
        trends = analysis.get("trendAnalysis", {}).get("keyTrends", [])
        for trend in trends:
            title = trend.get("title", "")
            if title in trend_titles_seen:
                continue
            trend_titles_seen.add(title)
            p(f"▶ {title}", size=12, bold=True, color=COLOR_COMP)
            desc = trend.get("description", "")
            if desc:
                p(desc, size=11)
            opp = trend.get("businessOpportunity", trend.get("strategicResponse", ""))
            if opp:
                p(f"사업 기회: {opp}", size=11, indent=0.5)
            reg = trend.get("regulatoryConsideration", "")
            if reg:
                p(f"규제 고려: {reg}", size=11, indent=0.5)
            p(f"(출처: {role} 분석)", size=9, color=COLOR_META)
            doc.add_paragraph()

    # ── 섹션 3: AI MSP 사업 추진 전략 ────────────────────
    p("3. AI MSP 사업 추진 전략", size=15, bold=True, color=COLOR_SECTION)
    doc.add_paragraph()

    priority_order = {"HIGH": 0, "MEDIUM": 1, "LOW": 2}
    all_strategies = []
    for analysis in analyses:
        role       = analysis.get("expertRole", "")
        strategies = analysis.get("strategyRecommendations", [])
        for s in strategies:
            s["_expertRole"] = role
            all_strategies.append(s)

    all_strategies.sort(key=lambda s: priority_order.get(s.get("priority", "LOW"), 99))

    for strategy in all_strategies:
        priority = strategy.get("priority", "MEDIUM")
        label = {"HIGH": "[긴급]", "MEDIUM": "[중요]", "LOW": "[검토]"}.get(priority, "[중요]")
        role  = strategy.get("_expertRole", "")
        title = strategy.get("title", "")
        p(f"{label} {title}  ({role})", size=12, bold=True, color=COLOR_COMP)

        rationale = strategy.get("rationale", "")
        if rationale:
            p(f"근거: {rationale}", size=11)

        diff = strategy.get("competitiveDifferentiation", "")
        if diff:
            p(f"차별화: {diff}", size=11, indent=0.5)

        target = strategy.get("targetCustomer", "")
        if target:
            p(f"타겟: {target}", size=11, indent=0.5)

        p("실행 액션:", size=11, bold=True)
        for action in strategy.get("actions", []):
            bullet(action)

        outcome = strategy.get("expectedOutcome", "")
        if outcome:
            p(f"기대 효과: {outcome}", size=11, indent=0.5)

        risk = strategy.get("riskConsideration", "")
        if risk:
            p(f"리스크: {risk}", size=11, indent=0.5)

        doc.add_paragraph()

    # ── 푸터 ─────────────────────────────────────────
    p("— AI MSP Intelligence Platform | 전문가 에이전트 팀 생성 레포트 —",
      size=9, color=COLOR_FOOTER, align=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


def main():
    if len(sys.argv) < 2:
        # 인자 없으면 가장 최근 workspace 사용
        workspaces = sorted(
            (PROJECT_ROOT / "_workspace").glob("weekly_*"), reverse=True)
        if not workspaces:
            print("❌ workspace 없음. fetch-weekly-data.py 먼저 실행하세요.", file=sys.stderr)
            sys.exit(1)
        workspace = workspaces[0]
    else:
        workspace = Path(sys.argv[1])

    print(f"📂 워크스페이스: {workspace}")

    meta          = load_json(workspace / "meta.json")
    ai_msp        = load_json(workspace / "ai_msp_analysis.json")
    ito           = load_json(workspace / "ito_analysis.json")
    it_strategy   = load_json(workspace / "it_strategy_analysis.json")

    analyses = [a for a in [ai_msp, ito, it_strategy] if a]
    if not analyses:
        print("❌ 전문가 분석 파일이 하나도 없습니다.", file=sys.stderr)
        sys.exit(1)

    print(f"✅ 분석 파일 로드: {[a.get('expertRole', '?') for a in analyses]}")

    week_start = meta.get("weekStart", workspace.name.replace("weekly_", ""))

    try:
        doc = generate_word(workspace, meta, analyses)
        REPORT_DIR.mkdir(parents=True, exist_ok=True)
        filename    = f"{week_start}_ai-msp-weekly-report.docx"
        output_path = REPORT_DIR / filename
        doc.save(str(output_path))
        print(f"\n✅ 레포트 생성 완료: {output_path.absolute()}")
        print(f"   파일 크기: {output_path.stat().st_size:,} bytes")
        print(f"REPORT_PATH={output_path.absolute()}")

        # 운영 서버에 업로드 → 원격 다운로드 URL 확보
        print("\n☁️  운영 서버 업로드 중...")
        download_url = upload_to_server(output_path, meta)
        if download_url:
            print(f"✅ 업로드 완료!")
            print(f"🔗 다운로드 URL: {download_url}")
        else:
            print("⚠️  업로드 실패 — 로컬 파일은 정상 생성됨")

    except ImportError:
        print("❌ python-docx 미설치: pip3 install python-docx", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
